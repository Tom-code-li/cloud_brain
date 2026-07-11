"""
将 模块材料包-患者系统.md 转换为 Word 文档 (.docx)

使用方法（在 cmd.exe 中执行）：
  pip install python-docx
  python md2docx.py
"""

import re
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)


def parse_markdown(filepath):
    """简易 Markdown 解析，按行处理"""
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
    return lines


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph_styled(doc, text, bold=False, italic=False, font_size=None, color=None, alignment=None):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_table_from_lines(doc, lines, start_idx, end_idx):
    """解析类表格行（| 分隔）为 Word 表格"""
    rows = []
    col_count = 0
    i = start_idx
    while i < end_idx:
        line = lines[i].strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            # 跳过分隔行 (|---|)
            if re.match(r"^[\s\-:]+$", "".join(cells)):
                i += 1
                continue
            if cells:
                col_count = max(col_count, len(cells))
                rows.append(cells)
        i += 1

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            cell.text = row_data[col_idx] if col_idx < len(row_data) else ""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()  # 表后空行


def convert_md_to_docx(md_path, docx_path):
    lines = parse_markdown(md_path)
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    i = 0
    in_code_block = False
    in_table = False
    table_start = -1

    while i < len(lines):
        line = lines[i].rstrip()

        # 跳过 frontmatter (---)
        if line.strip() == "---" and i < 3:
            while i < len(lines):
                i += 1
                if i < len(lines) and lines[i].strip() == "---":
                    i += 1
                    break
            continue

        # 代码块
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # 表格行
        if line.startswith("|") and line.endswith("|"):
            if not in_table:
                in_table = True
                table_start = i
            i += 1
            continue
        else:
            if in_table:
                add_table_from_lines(doc, lines, table_start, i)
                in_table = False
                table_start = -1
                # 跳过空行
                while i < len(lines) and lines[i].strip() == "":
                    i += 1
                continue

        # 空行
        if line.strip() == "":
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_heading_styled(doc, text, level=level)
            i += 1
            continue

        # 无序列表
        list_match = re.match(r"^[\*\-\+]\s+(.+)$", line)
        if list_match:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(list_match.group(1).strip())
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(11)
            i += 1
            continue

        # 有序列表
        olist_match = re.match(r"^\d+[\.\）\\)]\s+(.+)$", line)
        if olist_match:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(olist_match.group(1).strip())
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(11)
            i += 1
            continue

        # 普通段落（处理加粗、斜体、行内代码）
        text = line
        p = doc.add_paragraph()
        # 处理 **bold**
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            run.font.size = Pt(11)

        i += 1

    # 处理文件末尾的表格
    if in_table:
        add_table_from_lines(doc, lines, table_start, len(lines))

    doc.save(docx_path)
    print(f"✓ 转换成功: {docx_path}")


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, "模块材料包-患者系统.md")
    docx_file = os.path.join(script_dir, "模块材料包-患者系统.docx")

    if not os.path.exists(md_file):
        print(f"✗ 找不到 Markdown 文件: {md_file}")
        exit(1)

    convert_md_to_docx(md_file, docx_file)
